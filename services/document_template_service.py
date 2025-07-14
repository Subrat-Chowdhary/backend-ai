"""
Document Template Service for Light Version
Handles Word document generation for resume exports
"""
import logging
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import io
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentTemplateService:
    def __init__(self):
        self.templates = {
            'professional': self._create_professional_template,
            'modern': self._create_modern_template,
            'compact': self._create_compact_template
        }
    
    def generate_resume_document(self, resumes: List[Dict], template_name: str = 'professional') -> bytes:
        """Generate Word document with selected resumes"""
        try:
            if template_name not in self.templates:
                template_name = 'professional'
            
            # Create document using selected template
            doc = self.templates[template_name](resumes)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.getvalue()
        
        except Exception as e:
            logger.error(f"Document generation failed: {e}")
            raise Exception(f"Failed to generate document: {str(e)}")
    
    def _create_professional_template(self, resumes: List[Dict]) -> Document:
        """Create professional template document"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add title
        title = doc.add_heading('Selected Resume Profiles', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation info
        info_para = doc.add_paragraph()
        info_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        info_para.add_run(f"\nTotal Profiles: {len(resumes)}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
        
        # Process each resume
        for idx, resume in enumerate(resumes, 1):
            self._add_professional_resume(doc, resume, idx)
            
            # Add page break between resumes (except for last one)
            if idx < len(resumes):
                doc.add_page_break()
        
        return doc
    
    def _add_professional_resume(self, doc: Document, resume: Dict, index: int):
        """Add a single resume in professional format"""
        
        # Resume header
        header = doc.add_heading(f"Profile #{index}", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Personal Information Section
        doc.add_heading('Personal Information', level=2)
        
        # Create a table for personal info
        personal_table = doc.add_table(rows=0, cols=2)
        personal_table.style = 'Table Grid'
        
        # Add personal information rows
        personal_info = [
            ('Name', resume.get('name', 'Not Available')),
            ('Email', resume.get('email_id', 'Not Available')),
            ('Phone', resume.get('phone_number', 'Not Available')),
            ('Location', resume.get('location', 'Not Available')),
            ('Current Job Title', resume.get('current_job_title', 'Not Available')),
            ('LinkedIn', resume.get('linkedin_url', 'Not Available')),
            ('GitHub', resume.get('github_url', 'Not Available')),
            ('Similarity Score', f"{resume.get('similarity_score', 0):.2%}"),
        ]
        
        for label, value in personal_info:
            row = personal_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            # Make label bold
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Professional Summary
        if resume.get('objective'):
            doc.add_heading('Professional Summary', level=2)
            doc.add_paragraph(resume.get('objective', ''))
        
        # Skills Section
        if resume.get('skills'):
            doc.add_heading('Skills', level=2)
            skills = resume.get('skills', [])
            if isinstance(skills, list):
                skills_text = ', '.join(skills)
            else:
                skills_text = str(skills)
            doc.add_paragraph(skills_text)
        
        # Experience Section
        if resume.get('experience_summary'):
            doc.add_heading('Experience Summary', level=2)
            doc.add_paragraph(resume.get('experience_summary', ''))
        
        # Companies Worked With
        if resume.get('companies_worked_with_duration'):
            doc.add_heading('Work History', level=2)
            companies = resume.get('companies_worked_with_duration', [])
            if isinstance(companies, list):
                for company in companies:
                    doc.add_paragraph(f"• {company}", style='List Bullet')
            else:
                doc.add_paragraph(str(companies))
        
        # Education/Qualifications
        if resume.get('qualifications_summary'):
            doc.add_heading('Qualifications', level=2)
            doc.add_paragraph(resume.get('qualifications_summary', ''))
        
        # Projects
        if resume.get('projects'):
            doc.add_heading('Projects', level=2)
            projects = resume.get('projects', [])
            if isinstance(projects, list):
                for project in projects:
                    doc.add_paragraph(f"• {project}", style='List Bullet')
            else:
                doc.add_paragraph(str(projects))
        
        # Additional Information
        additional_info = []
        if resume.get('availability_status'):
            additional_info.append(f"Availability: {resume.get('availability_status')}")
        if resume.get('work_authorization_status'):
            additional_info.append(f"Work Authorization: {resume.get('work_authorization_status')}")
        if resume.get('_original_filename'):
            additional_info.append(f"Source File: {resume.get('_original_filename')}")
        
        if additional_info:
            doc.add_heading('Additional Information', level=2)
            for info in additional_info:
                doc.add_paragraph(f"• {info}", style='List Bullet')
    
    def _create_modern_template(self, resumes: List[Dict]) -> Document:
        """Create modern template document"""
        doc = Document()
        
        # Title with modern styling
        title = doc.add_heading('Resume Portfolio', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add summary table
        summary_table = doc.add_table(rows=1, cols=3)
        summary_table.style = 'Light Shading Accent 1'
        
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = 'Total Profiles'
        header_cells[1].text = 'Generated Date'
        header_cells[2].text = 'Template'
        
        data_row = summary_table.add_row()
        data_cells = data_row.cells
        data_cells[0].text = str(len(resumes))
        data_cells[1].text = datetime.now().strftime('%Y-%m-%d')
        data_cells[2].text = 'Modern'
        
        doc.add_page_break()
        
        # Process each resume with modern styling
        for idx, resume in enumerate(resumes, 1):
            self._add_modern_resume(doc, resume, idx)
            if idx < len(resumes):
                doc.add_page_break()
        
        return doc
    
    def _add_modern_resume(self, doc: Document, resume: Dict, index: int):
        """Add resume in modern format"""
        # Modern header with colored background
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(f"CANDIDATE #{index:02d}")
        header_run.bold = True
        header_run.font.size = Pt(16)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Name in large font
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(resume.get('name', 'Name Not Available').upper())
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info in one line
        contact_para = doc.add_paragraph()
        contact_info = []
        if resume.get('email_id', 'Not Available') != 'Not Available':
            contact_info.append(resume.get('email_id'))
        if resume.get('phone_number', 'Not Available') != 'Not Available':
            contact_info.append(resume.get('phone_number'))
        if resume.get('location', 'Not Available') != 'Not Available':
            contact_info.append(resume.get('location'))
        
        contact_para.add_run(' | '.join(contact_info))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Match score prominently displayed
        score_para = doc.add_paragraph()
        score_run = score_para.add_run(f"MATCH SCORE: {resume.get('similarity_score', 0):.1%}")
        score_run.bold = True
        score_run.font.size = Pt(14)
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space
        
        # Two-column layout for content
        content_table = doc.add_table(rows=1, cols=2)
        content_table.autofit = False
        content_table.columns[0].width = Inches(3)
        content_table.columns[1].width = Inches(3.5)
        
        left_cell = content_table.cell(0, 0)
        right_cell = content_table.cell(0, 1)
        
        # Left column content
        left_para = left_cell.paragraphs[0]
        left_para.add_run("SKILLS").bold = True
        left_cell.add_paragraph()
        
        skills = resume.get('skills', [])
        if isinstance(skills, list):
            for skill in skills[:10]:  # Limit to top 10 skills
                left_cell.add_paragraph(f"• {skill}")
        
        left_cell.add_paragraph()
        left_para = left_cell.add_paragraph()
        left_para.add_run("CONTACT").bold = True
        
        if resume.get('linkedin_url', 'Not Available') != 'Not Available':
            left_cell.add_paragraph(f"LinkedIn: {resume.get('linkedin_url')}")
        if resume.get('github_url', 'Not Available') != 'Not Available':
            left_cell.add_paragraph(f"GitHub: {resume.get('github_url')}")
        
        # Right column content
        right_para = right_cell.paragraphs[0]
        right_para.add_run("EXPERIENCE").bold = True
        right_cell.add_paragraph()
        
        if resume.get('current_job_title'):
            right_cell.add_paragraph(f"Current Role: {resume.get('current_job_title')}")
        
        if resume.get('experience_summary'):
            exp_text = resume.get('experience_summary')
            if len(exp_text) > 300:
                exp_text = exp_text[:300] + "..."
            right_cell.add_paragraph(exp_text)
        
        if resume.get('companies_worked_with_duration'):
            right_cell.add_paragraph()
            right_para = right_cell.add_paragraph()
            right_para.add_run("WORK HISTORY").bold = True
            
            companies = resume.get('companies_worked_with_duration', [])
            if isinstance(companies, list):
                for company in companies[:5]:  # Limit to 5 companies
                    right_cell.add_paragraph(f"• {company}")
    
    def _create_compact_template(self, resumes: List[Dict]) -> Document:
        """Create compact template document - multiple resumes per page"""
        doc = Document()
        
        # Set narrow margins for compact layout
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Title
        title = doc.add_heading('Resume Summary Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary info
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Total Profiles: {len(resumes)}")
        doc.add_paragraph()
        
        # Process resumes in compact format
        for idx, resume in enumerate(resumes, 1):
            self._add_compact_resume(doc, resume, idx)
            
            # Add separator between resumes
            if idx < len(resumes):
                doc.add_paragraph("─" * 80)
    
    def _add_compact_resume(self, doc: Document, resume: Dict, index: int):
        """Add resume in compact format"""
        # Compact header
        header_para = doc.add_paragraph()
        header_para.add_run(f"{index}. {resume.get('name', 'N/A')} | {resume.get('current_job_title', 'N/A')} | Score: {resume.get('similarity_score', 0):.1%}").bold = True
        
        # Contact in one line
        contact_para = doc.add_paragraph()
        contact_info = [
            resume.get('email_id', 'N/A'),
            resume.get('phone_number', 'N/A'),
            resume.get('location', 'N/A')
        ]
        contact_para.add_run(f"Contact: {' | '.join([info for info in contact_info if info != 'N/A'])}")
        
        # Skills in one line
        skills = resume.get('skills', [])
        if skills:
            skills_text = ', '.join(skills[:8]) if isinstance(skills, list) else str(skills)[:100]
            doc.add_paragraph(f"Skills: {skills_text}")
        
        # Brief experience
        if resume.get('experience_summary'):
            exp_text = resume.get('experience_summary')
            if len(exp_text) > 150:
                exp_text = exp_text[:150] + "..."
            doc.add_paragraph(f"Experience: {exp_text}")
        
        doc.add_paragraph()  # Space between entries

    def get_available_templates(self) -> List[str]:
        """Get list of available templates"""
        return list(self.templates.keys())

# Create global instance
document_template_service = DocumentTemplateService()